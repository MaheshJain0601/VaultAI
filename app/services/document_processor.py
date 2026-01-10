"""Document processing service using LangChain for text extraction and chunking."""
import re
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass

import pdfplumber
from docx import Document as DocxDocument

# LangChain imports for text splitting
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

from app.config import settings


@dataclass
class ExtractedText:
    """Extracted text with metadata."""
    content: str
    page_count: int
    word_count: int
    character_count: int
    pages: List[Dict]  # List of {page_number, content}
    metadata: Dict


@dataclass
class TextChunk:
    """A chunk of text with position metadata."""
    content: str
    chunk_index: int
    page_number: Optional[int]
    start_char: int
    end_char: int
    token_count: int


class DocumentProcessor:
    """
    Service for processing and extracting text from documents.
    
    Uses LangChain's RecursiveCharacterTextSplitter for intelligent chunking
    that respects semantic boundaries (paragraphs, sentences, words).
    
    LangChain components used:
    - RecursiveCharacterTextSplitter: Smart text chunking with overlap
    - Document: LangChain document schema for metadata handling
    """
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        
        # Initialize LangChain text splitter
        # RecursiveCharacterTextSplitter tries to split on these separators in order,
        # keeping semantically related text together
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=[
                "\n\n",  # Paragraphs
                "\n",    # Lines
                ". ",    # Sentences
                "! ",    # Exclamation sentences
                "? ",    # Question sentences
                "; ",    # Semicolon clauses
                ", ",    # Comma clauses
                " ",     # Words
                ""       # Characters (last resort)
            ],
            is_separator_regex=False
        )
    
    def extract_text(self, file_path: str, file_type: str) -> ExtractedText:
        """
        Extract text from a document.
        
        Args:
            file_path: Path to the document file
            file_type: Type of document (pdf, docx, txt, md)
            
        Returns:
            ExtractedText object with content and metadata
        """
        extractors = {
            "pdf": self._extract_from_pdf,
            "docx": self._extract_from_docx,
            "txt": self._extract_from_text,
            "md": self._extract_from_text,
        }
        
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")
            
        return extractor(file_path)
    
    def _extract_from_pdf(self, file_path: str) -> ExtractedText:
        """Extract text from PDF using pdfplumber."""
        pages = []
        full_text = []
        metadata = {}
        
        with pdfplumber.open(file_path) as pdf:
            metadata = {
                "author": pdf.metadata.get("Author", ""),
                "title": pdf.metadata.get("Title", ""),
                "subject": pdf.metadata.get("Subject", ""),
                "creator": pdf.metadata.get("Creator", ""),
            }
            
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                pages.append({
                    "page_number": i + 1,
                    "content": page_text
                })
                full_text.append(page_text)
        
        content = "\n\n".join(full_text)
        
        return ExtractedText(
            content=content,
            page_count=len(pages),
            word_count=len(content.split()),
            character_count=len(content),
            pages=pages,
            metadata=metadata
        )
    
    def _extract_from_docx(self, file_path: str) -> ExtractedText:
        """Extract text from DOCX."""
        doc = DocxDocument(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        content = "\n\n".join(paragraphs)
        
        # Extract core properties
        metadata = {}
        if doc.core_properties:
            metadata = {
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or "",
                "subject": doc.core_properties.subject or "",
            }
        
        return ExtractedText(
            content=content,
            page_count=1,  # DOCX doesn't have clear page boundaries
            word_count=len(content.split()),
            character_count=len(content),
            pages=[{"page_number": 1, "content": content}],
            metadata=metadata
        )
    
    def _extract_from_text(self, file_path: str) -> ExtractedText:
        """Extract text from TXT or MD files."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        return ExtractedText(
            content=content,
            page_count=1,
            word_count=len(content.split()),
            character_count=len(content),
            pages=[{"page_number": 1, "content": content}],
            metadata={}
        )
    
    def chunk_text(
        self, 
        extracted: ExtractedText,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> List[TextChunk]:
        """
        Split text into overlapping chunks using LangChain's RecursiveCharacterTextSplitter.
        
        This method uses LangChain for intelligent chunking that:
        1. Respects semantic boundaries (paragraphs, sentences)
        2. Maintains configurable overlap for context preservation
        3. Tracks page numbers for citation purposes
        
        Args:
            extracted: ExtractedText from document
            chunk_size: Optional override for chunk size
            chunk_overlap: Optional override for overlap
            
        Returns:
            List of TextChunk objects with metadata
        """
        # Create custom splitter if sizes are overridden
        if chunk_size or chunk_overlap:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size or self.chunk_size,
                chunk_overlap=chunk_overlap or self.chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""]
            )
        else:
            splitter = self.text_splitter
        
        chunks = []
        chunk_index = 0
        
        # Process each page separately to maintain page references
        for page_info in extracted.pages:
            page_number = page_info["page_number"]
            page_content = page_info["content"]
            
            if not page_content.strip():
                continue
            
            # Use LangChain to create documents with metadata
            page_doc = LangChainDocument(
                page_content=page_content,
                metadata={"page_number": page_number}
            )
            
            # Split using LangChain's text splitter
            split_docs = splitter.split_documents([page_doc])
            
            # Convert LangChain documents to our TextChunk format
            current_pos = 0
            for doc in split_docs:
                content = doc.page_content
                
                # Find position in original text
                start_char = page_content.find(content[:50], current_pos)
                if start_char == -1:
                    start_char = current_pos
                end_char = start_char + len(content)
                current_pos = end_char
                
                chunks.append(TextChunk(
                    content=content,
                    chunk_index=chunk_index,
                    page_number=doc.metadata.get("page_number", page_number),
                    start_char=start_char,
                    end_char=end_char,
                    token_count=self._estimate_tokens(content)
                ))
                chunk_index += 1
        
        return chunks
    
    def chunk_text_langchain(
        self,
        extracted: ExtractedText
    ) -> List[LangChainDocument]:
        """
        Split text into LangChain Document objects directly.
        
        Useful when you want to use LangChain's vector stores or
        other components that expect LangChain Document format.
        
        Returns:
            List of LangChain Document objects with metadata
        """
        all_docs = []
        
        for page_info in extracted.pages:
            page_number = page_info["page_number"]
            page_content = page_info["content"]
            
            if not page_content.strip():
                continue
            
            # Create LangChain document
            page_doc = LangChainDocument(
                page_content=page_content,
                metadata={
                    "page_number": page_number,
                    "source": f"page_{page_number}"
                }
            )
            
            # Split using LangChain
            split_docs = self.text_splitter.split_documents([page_doc])
            
            # Add chunk indices
            for i, doc in enumerate(split_docs):
                doc.metadata["chunk_index"] = len(all_docs) + i
            
            all_docs.extend(split_docs)
        
        return all_docs
    
    def _estimate_tokens(self, text: str) -> int:
        """
        Estimate token count.
        
        Uses a simple approximation (~4 characters per token for English).
        For more accurate counts, use tiktoken directly.
        """
        return len(text) // 4


# Singleton instance
document_processor = DocumentProcessor()
